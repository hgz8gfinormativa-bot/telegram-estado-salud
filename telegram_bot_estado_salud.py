import html
import json
import logging
import os
import re
import tempfile
import asyncio
from datetime import datetime
from typing import Any, Dict, List

import fitz
from openai import OpenAI
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from telegram import (
    Update,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
)
from telegram.ext import (
    Application,
    CommandHandler,
    ContextTypes,
    MessageHandler,
    CallbackQueryHandler,
    filters,
)
from docx import Document
from docx.shared import Pt

# OCR opcional
OCR_ENABLED = os.getenv("OCR_ENABLED", "false").lower() == "true"
try:
    import pytesseract
    from PIL import Image
    OCR_LIBS_AVAILABLE = True
except Exception:
    OCR_LIBS_AVAILABLE = False

# =========================
# CONFIGURACIÓN
# =========================

logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.1-8b-instant")

INSTITUTION_NAME = "HGZ/MF 8 DR. GILBERTO FLORES IZQUIERDO"

INCLUDE_PDF_FILE = os.getenv("INCLUDE_PDF_FILE", "true").lower() == "true"
INCLUDE_DOCX_FILE = os.getenv("INCLUDE_DOCX_FILE", "true").lower() == "true"

MAX_CHUNK_CHARS = int(os.getenv("MAX_CHUNK_CHARS", "12000"))
MAX_PROMPT_CHARS = int(os.getenv("MAX_PROMPT_CHARS", "50000"))
MAX_EXTENDED_CONTEXT_CHARS = int(os.getenv("MAX_EXTENDED_CONTEXT_CHARS", "45000"))

TELEGRAM_MAX_MESSAGE = 4000
MAX_FILE_SIZE_MB = int(os.getenv("MAX_FILE_SIZE_MB", "20"))
MIN_TEXT_LENGTH = int(os.getenv("MIN_TEXT_LENGTH", "200"))

if not TELEGRAM_BOT_TOKEN:
    raise RuntimeError("Falta TELEGRAM_BOT_TOKEN")

if not GROQ_API_KEY:
    raise RuntimeError("Falta GROQ_API_KEY")

client = OpenAI(
    api_key=GROQ_API_KEY,
    base_url="https://api.groq.com/openai/v1",
)

# =========================
# PROMPTS
# =========================

SYSTEM_PROMPT_EXTRACT = """
Eres un asistente clínico-administrativo especializado en análisis documental de notas médicas.

Tu tarea es EXTRAER información documentada de un fragmento de notas médicas.

Reglas obligatorias:
- No inventes información.
- No infieras diagnósticos nuevos.
- No completes datos faltantes.
- Si un dato no aparece con claridad, usa "No se documenta".
- No hagas recomendaciones.
- Debes devolver JSON válido.
- La redacción debe ser formal, objetiva y clínica.
- Si hay fechas, consérvalas como aparezcan en el documento.
- Si existen signos vitales, consérvalos textualmente.
- Si existe nombre del médico tratante, consérvalo.
- Si existe servicio clínico, consérvalo.
- Si existen medicamentos, procedimientos, interconsultas o estudios, consérvalos textualmente cuando sea posible.

Devuelve exactamente esta estructura JSON:
{
  "paciente": {
    "nombre": "",
    "nss": "",
    "edad": "",
    "sexo": ""
  },
  "fechas_relevantes": [],
  "fuentes": [],
  "medicos_tratantes": [],
  "servicios_clinicos": [],
  "signos_vitales": [],
  "notas_cronologicas": [
    {
      "fecha": "",
      "servicio": "",
      "tipo_nota": "",
      "medico_tratante": "",
      "signos_vitales": [],
      "resumen": "",
      "diagnosticos": [],
      "tratamiento": [],
      "estado_actual": "",
      "pronostico": ""
    }
  ],
  "diagnosticos_globales": [],
  "tratamiento_global_documentado": [],
  "estado_global_documentado": "",
  "pronostico_global_documentado": "",
  "observaciones": []
}
"""

SYSTEM_PROMPT_FINAL_ESTADO = """
Eres un asistente clínico-administrativo.

Tu tarea es elaborar un BORRADOR de ESTADO DE SALUD basado únicamente en información previamente consolidada.

Reglas obligatorias:
- No inventes información.
- No agregues diagnósticos nuevos.
- No hagas recomendaciones médicas.
- Si algo no está documentado, escribe "No se documenta".
- Redacción formal, breve, clara e institucional.
- Devuelve JSON válido.

Devuelve exactamente esta estructura JSON:
{
  "tipo_documento": "estado_salud",
  "institucion": "",
  "identificacion": {
    "nombre": "",
    "nss": "",
    "edad": "",
    "sexo": ""
  },
  "fecha_referencia": "",
  "fuentes": [],
  "medicos_tratantes": [],
  "servicios_clinicos": [],
  "signos_vitales": [],
  "resumen_clinico": "",
  "diagnosticos_documentados": [],
  "estado_actual": "",
  "tratamiento_actual_documentado": [],
  "pronostico_documentado": "",
  "observaciones": [],
  "texto_final": ""
}
"""

SYSTEM_PROMPT_FINAL_RESUMEN = """
Eres un asistente clínico-administrativo.

Tu tarea es elaborar un RESUMEN CLÍNICO DOCUMENTAL basado únicamente en información previamente consolidada y ampliada.

Reglas obligatorias:
- No inventes información.
- No agregues datos nuevos.
- Si algo no está documentado, escribe "No se documenta".
- Redacción formal, objetiva y ordenada.
- Procura incorporar la mayor cantidad posible de información documental relevante.
- Devuelve JSON válido.

Devuelve exactamente esta estructura JSON:
{
  "tipo_documento": "resumen_clinico",
  "institucion": "",
  "identificacion": {
    "nombre": "",
    "nss": "",
    "edad": "",
    "sexo": ""
  },
  "fecha_referencia": "",
  "fuentes": [],
  "medicos_tratantes": [],
  "servicios_clinicos": [],
  "signos_vitales": [],
  "resumen_clinico": "",
  "diagnosticos_documentados": [],
  "tratamiento_actual_documentado": [],
  "observaciones": [],
  "texto_final": ""
}
"""

SYSTEM_PROMPT_FINAL_CRONOLOGIA = """
Eres un asistente clínico-administrativo.

Tu tarea es elaborar una CRONOLOGÍA MÉDICA DOCUMENTAL basada únicamente en información previamente consolidada y ampliada.

Reglas obligatorias:
- No inventes información.
- No agregues datos nuevos.
- Ordena cronológicamente en la medida de lo posible con base en las fechas documentadas.
- Si una fecha no está documentada claramente, conserva "No se documenta".
- Procura conservar el mayor detalle documental útil.
- Devuelve JSON válido.

Devuelve exactamente esta estructura JSON:
{
  "tipo_documento": "cronologia_medica",
  "institucion": "",
  "identificacion": {
    "nombre": "",
    "nss": "",
    "edad": "",
    "sexo": ""
  },
  "fecha_referencia": "",
  "fuentes": [],
  "medicos_tratantes": [],
  "servicios_clinicos": [],
  "signos_vitales": [],
  "cronologia": [
    {
      "fecha": "",
      "servicio": "",
      "tipo_nota": "",
      "medico_tratante": "",
      "signos_vitales": [],
      "resumen": "",
      "diagnosticos": [],
      "tratamiento": [],
      "estado_actual": "",
      "pronostico": ""
    }
  ],
  "texto_final": ""
}
"""

SYSTEM_PROMPT_FINAL_ESTADO_FAMILIAR = """
Eres un asistente clínico-administrativo.

Tu tarea es elaborar un BORRADOR DE ESTADO DE SALUD PARA FAMILIAR basado únicamente en información previamente consolidada.

Reglas obligatorias:
- No inventes información.
- No agregues diagnósticos nuevos.
- No hagas recomendaciones médicas.
- Usa lenguaje claro, sencillo y respetuoso.
- No sustituyas valoración médica.
- Si algo no está documentado, escribe "No se documenta".
- Devuelve JSON válido.

Devuelve exactamente esta estructura JSON:
{
  "tipo_documento": "estado_familiar",
  "institucion": "",
  "identificacion": {
    "nombre": "",
    "nss": "",
    "edad": "",
    "sexo": ""
  },
  "fecha_referencia": "",
  "fuentes": [],
  "servicios_clinicos": [],
  "resumen_clinico": "",
  "estado_actual": "",
  "tratamiento_actual_documentado": [],
  "observaciones": [],
  "texto_final": ""
}
"""

SYSTEM_PROMPT_FINAL_ESTADO_AUTORIDAD = """
Eres un asistente clínico-administrativo.

Tu tarea es elaborar un BORRADOR DE ESTADO DE SALUD PARA AUTORIDAD basado únicamente en información previamente consolidada y ampliada.

Reglas obligatorias:
- No inventes información.
- No agregues diagnósticos nuevos.
- No emitas opiniones médico-legales.
- Usa redacción formal, técnica, objetiva e institucional.
- Si algo no está documentado, escribe "No se documenta".
- Procura incorporar el mayor detalle documental relevante.
- Devuelve JSON válido.

Devuelve exactamente esta estructura JSON:
{
  "tipo_documento": "estado_autoridad",
  "institucion": "",
  "identificacion": {
    "nombre": "",
    "nss": "",
    "edad": "",
    "sexo": ""
  },
  "fecha_referencia": "",
  "fuentes": [],
  "medicos_tratantes": [],
  "servicios_clinicos": [],
  "signos_vitales": [],
  "resumen_clinico": "",
  "diagnosticos_documentados": [],
  "estado_actual": "",
  "tratamiento_actual_documentado": [],
  "pronostico_documentado": "",
  "observaciones": [],
  "texto_final": ""
}
"""

SYSTEM_PROMPT_FINAL_ESTADO_INSTITUCIONAL = """
Eres un asistente clínico-administrativo.

Tu tarea es elaborar un BORRADOR DE ESTADO DE SALUD INSTITUCIONAL basado únicamente en información previamente consolidada y ampliada.

Reglas obligatorias:
- No inventes información.
- No agregues diagnósticos nuevos.
- No emitas opinión pericial.
- Usa redacción formal, técnica, objetiva y administrativa.
- Si algo no está documentado, escribe "No se documenta".
- Debe ser útil como borrador institucional.
- Procura incorporar el mayor detalle documental relevante.
- Devuelve JSON válido.

Devuelve exactamente esta estructura JSON:
{
  "tipo_documento": "estado_institucional",
  "institucion": "",
  "identificacion": {
    "nombre": "",
    "nss": "",
    "edad": "",
    "sexo": ""
  },
  "fecha_referencia": "",
  "fuentes": [],
  "medicos_tratantes": [],
  "servicios_clinicos": [],
  "signos_vitales": [],
  "resumen_clinico": "",
  "diagnosticos_documentados": [],
  "estado_actual": "",
  "tratamiento_actual_documentado": [],
  "pronostico_documentado": "",
  "observaciones": [],
  "texto_final": ""
}
"""

# =========================
# UTILIDADES
# =========================

HELP_TEXT = (
    "Envíame un PDF con notas médicas usando alguno de estos comandos:\n\n"
    "/estado - Genera estado de salud\n"
    "/resumen - Genera resumen clínico\n"
    "/cronologia - Genera cronología médica\n"
    "/estado_familiar - Genera versión para familiar\n"
    "/estado_autoridad - Genera versión para autoridad\n"
    "/estado_institucional - Genera versión institucional\n"
    "/help - Muestra esta ayuda\n\n"
    "Después de enviar el comando, adjunta el PDF."
)

def track_bot_message_id(context: ContextTypes.DEFAULT_TYPE, message_id: int) -> None:
    context.user_data.setdefault("bot_message_ids", []).append(message_id)

async def send_tracked_text(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
    text: str,
    reply_markup=None
):
    msg = await update.message.reply_text(text, reply_markup=reply_markup)
    track_bot_message_id(context, msg.message_id)
    return msg

async def send_tracked_text_from_query(
    query,
    context: ContextTypes.DEFAULT_TYPE,
    text: str,
    reply_markup=None
):
    msg = await query.message.reply_text(text, reply_markup=reply_markup)
    track_bot_message_id(context, msg.message_id)
    return msg

def build_finish_keyboard() -> InlineKeyboardMarkup:
    keyboard = [
        [
            InlineKeyboardButton("🆕 Nuevo reporte", callback_data="nuevo_reporte"),
            InlineKeyboardButton("🧹 Limpiar mensajes del bot", callback_data="limpiar_bot"),
        ],
        [
            InlineKeyboardButton("📋 Ver menú", callback_data="ver_menu"),
        ],
    ]
    return InlineKeyboardMarkup(keyboard)

def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages = []

    for i, page in enumerate(doc, start=1):
        text = page.get_text("text")
        if text and text.strip():
            pages.append(f"\n--- Página {i} ---\n{text.strip()}")

    needs_ocr = not pages
    ocr_texts = []

    if needs_ocr and OCR_ENABLED and OCR_LIBS_AVAILABLE:
        logger.info("PDF sin texto extraíble; iniciando OCR.")
        for i, page in enumerate(doc, start=1):
            try:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
                    img_path = tmp_img.name
                    pix.save(img_path)

                img = Image.open(img_path)
                text = pytesseract.image_to_string(img, lang="spa+eng")
                if text and text.strip():
                    ocr_texts.append(f"\n--- Página {i} OCR ---\n{text.strip()}")

                try:
                    os.remove(img_path)
                except Exception:
                    pass
            except Exception as e:
                logger.warning("OCR falló en página %s: %s", i, e)

    doc.close()

    if pages:
        return "\n".join(pages).strip()
    if ocr_texts:
        return "\n".join(ocr_texts).strip()
    return ""

def split_text(text: str, max_chars: int = MAX_CHUNK_CHARS) -> List[str]:
    if len(text) <= max_chars:
        return [text]

    chunks = []
    start = 0
    length = len(text)

    while start < length:
        end = min(start + max_chars, length)
        if end < length:
            cut = text.rfind("\n", start, end)
            if cut > start + 1000:
                end = cut

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start = end

    return chunks

def safe_json_loads(content: str) -> Dict[str, Any]:
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        start = content.find("{")
        end = content.rfind("}")
        if start != -1 and end != -1 and end > start:
            return json.loads(content[start:end + 1])
        raise

def chat_json(system_prompt: str, user_prompt: str) -> Dict[str, Any]:
    try:
        response = client.chat.completions.create(
            model=GROQ_MODEL,
            temperature=0.2,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt[:MAX_PROMPT_CHARS]},
            ],
        )
        content = response.choices[0].message.content
        return safe_json_loads(content)
    except Exception:
        logger.exception("Fallo en chat_json")
        raise

def normalize_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()

def normalize_list(value: Any) -> List[str]:
    if not value:
        return []
    if isinstance(value, list):
        out = []
        seen = set()
        for item in value:
            item = normalize_text(item)
            if item and item not in seen:
                out.append(item)
                seen.add(item)
        return out
    value = normalize_text(value)
    return [value] if value else []

def format_list(items: Any) -> str:
    values = normalize_list(items)
    if not values:
        return "No se documenta"
    return "\n".join(f"• {x}" for x in values)

def validate_nss(nss: str) -> str:
    value = re.sub(r"\D", "", normalize_text(nss))
    if len(value) in (10, 11):
        return value
    return "No se documenta"

def clean_name(name: str) -> str:
    value = normalize_text(name)
    if not value:
        return "No se documenta"
    value = re.sub(r"\s+", " ", value)
    if len(re.findall(r"[A-Za-zÁÉÍÓÚÑáéíóúñ]", value)) < 3:
        return "No se documenta"
    return value

def sort_cronologia(entries: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    def parse_date(value: str):
        value = normalize_text(value)
        if not value or value.lower() == "no se documenta":
            return (1, datetime.max)

        patterns = [
            "%Y-%m-%d",
            "%d/%m/%Y",
            "%d-%m-%Y",
            "%d.%m.%Y",
        ]

        for fmt in patterns:
            try:
                return (0, datetime.strptime(value[:10], fmt))
            except Exception:
                pass

        match = re.search(r"(\d{2})[/-](\d{2})[/-](\d{4})", value)
        if match:
            d, m, y = match.groups()
            try:
                return (0, datetime(int(y), int(m), int(d)))
            except Exception:
                pass

        return (1, datetime.max)

    return sorted(entries, key=lambda x: parse_date(x.get("fecha", "")))

def escape_pdf_text(text: str) -> str:
    return html.escape(text or "").replace("\n", "<br/>")

def build_extended_context(merged: Dict[str, Any], max_chars: int = MAX_EXTENDED_CONTEXT_CHARS) -> str:
    sections: List[str] = []

    sections.append("DATOS DEL PACIENTE:")
    sections.append(json.dumps(merged.get("paciente", {}), ensure_ascii=False, indent=2))

    sections.append("\nFECHAS RELEVANTES:")
    sections.append("\n".join(merged.get("fechas_relevantes", [])) or "No se documenta")

    sections.append("\nSERVICIOS CLÍNICOS:")
    sections.append("\n".join(merged.get("servicios_clinicos", [])) or "No se documenta")

    sections.append("\nMÉDICOS TRATANTES:")
    sections.append("\n".join(merged.get("medicos_tratantes", [])) or "No se documenta")

    sections.append("\nSIGNOS VITALES:")
    sections.append("\n".join(merged.get("signos_vitales", [])) or "No se documenta")

    sections.append("\nDIAGNÓSTICOS GLOBALES DOCUMENTADOS:")
    sections.append("\n".join(merged.get("diagnosticos_globales", [])) or "No se documenta")

    sections.append("\nTRATAMIENTO GLOBAL DOCUMENTADO:")
    sections.append("\n".join(merged.get("tratamiento_global_documentado", [])) or "No se documenta")

    sections.append("\nESTADO GLOBAL DOCUMENTADO:")
    sections.append("\n".join(merged.get("estado_global_documentado", [])) or "No se documenta")

    sections.append("\nPRONÓSTICO GLOBAL DOCUMENTADO:")
    sections.append("\n".join(merged.get("pronostico_global_documentado", [])) or "No se documenta")

    sections.append("\nOBSERVACIONES:")
    sections.append("\n".join(merged.get("observaciones", [])) or "No se documenta")

    sections.append("\nCRONOLOGÍA DOCUMENTADA:")
    cronologia = merged.get("notas_cronologicas", [])
    if not cronologia:
        sections.append("No se documenta")
    else:
        for i, item in enumerate(cronologia, start=1):
            sections.append(
                f"{i}. Fecha: {item.get('fecha', 'No se documenta')}\n"
                f"Servicio: {item.get('servicio', 'No se documenta')}\n"
                f"Tipo de nota: {item.get('tipo_nota', 'No se documenta')}\n"
                f"Médico tratante: {item.get('medico_tratante', 'No se documenta')}\n"
                f"Signos vitales: {', '.join(item.get('signos_vitales', [])) or 'No se documenta'}\n"
                f"Resumen: {item.get('resumen', 'No se documenta')}\n"
                f"Diagnósticos: {', '.join(item.get('diagnosticos', [])) or 'No se documenta'}\n"
                f"Tratamiento: {', '.join(item.get('tratamiento', [])) or 'No se documenta'}\n"
                f"Estado actual: {item.get('estado_actual', 'No se documenta')}\n"
                f"Pronóstico: {item.get('pronostico', 'No se documenta')}\n"
            )

    full_text = "\n\n".join(sections)
    return full_text[:max_chars]

# =========================
# EXTRACCIÓN Y CONSOLIDACIÓN
# =========================

def summarize_chunk(chunk_text: str, index: int, total: int) -> Dict[str, Any]:
    prompt = f"""
Analiza el fragmento {index} de {total} de un conjunto de notas médicas.

Texto:
{chunk_text}
"""
    return chat_json(SYSTEM_PROMPT_EXTRACT, prompt)

def merge_extractions(partials: List[Dict[str, Any]]) -> Dict[str, Any]:
    merged = {
        "paciente": {
            "nombre": "No se documenta",
            "nss": "No se documenta",
            "edad": "No se documenta",
            "sexo": "No se documenta",
        },
        "fechas_relevantes": [],
        "fuentes": [],
        "medicos_tratantes": [],
        "servicios_clinicos": [],
        "signos_vitales": [],
        "notas_cronologicas": [],
        "diagnosticos_globales": [],
        "tratamiento_global_documentado": [],
        "estado_global_documentado": [],
        "pronostico_global_documentado": [],
        "observaciones": [],
    }

    seen_fechas = set()
    seen_fuentes = set()
    seen_medicos = set()
    seen_servicios = set()
    seen_sv = set()
    seen_diag = set()
    seen_tx = set()
    seen_obs = set()
    seen_estado = set()
    seen_pron = set()
    seen_notes = set()

    for part in partials:
        paciente = part.get("paciente", {})
        nombre = clean_name(paciente.get("nombre"))
        if nombre != "No se documenta" and merged["paciente"]["nombre"] == "No se documenta":
            merged["paciente"]["nombre"] = nombre

        nss = validate_nss(paciente.get("nss", ""))
        if nss != "No se documenta" and merged["paciente"]["nss"] == "No se documenta":
            merged["paciente"]["nss"] = nss

        for key in ["edad", "sexo"]:
            val = normalize_text(paciente.get(key))
            if val and val.lower() != "no se documenta" and merged["paciente"][key] == "No se documenta":
                merged["paciente"][key] = val

        for fecha in normalize_list(part.get("fechas_relevantes")):
            if fecha not in seen_fechas:
                merged["fechas_relevantes"].append(fecha)
                seen_fechas.add(fecha)

        for fuente in normalize_list(part.get("fuentes")):
            if fuente not in seen_fuentes:
                merged["fuentes"].append(fuente)
                seen_fuentes.add(fuente)

        for medico in normalize_list(part.get("medicos_tratantes")):
            if medico not in seen_medicos:
                merged["medicos_tratantes"].append(medico)
                seen_medicos.add(medico)

        for servicio in normalize_list(part.get("servicios_clinicos")):
            if servicio not in seen_servicios:
                merged["servicios_clinicos"].append(servicio)
                seen_servicios.add(servicio)

        for sv in normalize_list(part.get("signos_vitales")):
            if sv not in seen_sv:
                merged["signos_vitales"].append(sv)
                seen_sv.add(sv)

        for diag in normalize_list(part.get("diagnosticos_globales")):
            if diag not in seen_diag:
                merged["diagnosticos_globales"].append(diag)
                seen_diag.add(diag)

        for tx in normalize_list(part.get("tratamiento_global_documentado")):
            if tx not in seen_tx:
                merged["tratamiento_global_documentado"].append(tx)
                seen_tx.add(tx)

        for obs in normalize_list(part.get("observaciones")):
            if obs not in seen_obs:
                merged["observaciones"].append(obs)
                seen_obs.add(obs)

        estado = normalize_text(part.get("estado_global_documentado"))
        if estado and estado.lower() != "no se documenta" and estado not in seen_estado:
            merged["estado_global_documentado"].append(estado)
            seen_estado.add(estado)

        pron = normalize_text(part.get("pronostico_global_documentado"))
        if pron and pron.lower() != "no se documenta" and pron not in seen_pron:
            merged["pronostico_global_documentado"].append(pron)
            seen_pron.add(pron)

        for note in part.get("notas_cronologicas", []):
            if not isinstance(note, dict):
                continue

            normalized_note = {
                "fecha": normalize_text(note.get("fecha")) or "No se documenta",
                "servicio": normalize_text(note.get("servicio")) or "No se documenta",
                "tipo_nota": normalize_text(note.get("tipo_nota")) or "No se documenta",
                "medico_tratante": normalize_text(note.get("medico_tratante")) or "No se documenta",
                "signos_vitales": normalize_list(note.get("signos_vitales")),
                "resumen": normalize_text(note.get("resumen")) or "No se documenta",
                "diagnosticos": normalize_list(note.get("diagnosticos")),
                "tratamiento": normalize_list(note.get("tratamiento")),
                "estado_actual": normalize_text(note.get("estado_actual")) or "No se documenta",
                "pronostico": normalize_text(note.get("pronostico")) or "No se documenta",
            }

            if normalized_note["medico_tratante"] != "No se documenta" and normalized_note["medico_tratante"] not in seen_medicos:
                merged["medicos_tratantes"].append(normalized_note["medico_tratante"])
                seen_medicos.add(normalized_note["medico_tratante"])

            if normalized_note["servicio"] != "No se documenta" and normalized_note["servicio"] not in seen_servicios:
                merged["servicios_clinicos"].append(normalized_note["servicio"])
                seen_servicios.add(normalized_note["servicio"])

            for sv in normalized_note["signos_vitales"]:
                if sv not in seen_sv:
                    merged["signos_vitales"].append(sv)
                    seen_sv.add(sv)

            note_key = json.dumps(normalized_note, ensure_ascii=False, sort_keys=True)
            if note_key not in seen_notes:
                merged["notas_cronologicas"].append(normalized_note)
                seen_notes.add(note_key)

    merged["notas_cronologicas"] = sort_cronologia(merged["notas_cronologicas"])
    return merged

def process_medical_text(text: str) -> Dict[str, Any]:
    chunks = split_text(text)
    partials = []

    for idx, chunk in enumerate(chunks, start=1):
        partial = summarize_chunk(chunk, idx, len(chunks))
        partials.append(partial)

    return merge_extractions(partials)

# =========================
# GENERADORES FINALES
# =========================

def build_estado_salud(merged: Dict[str, Any]) -> Dict[str, Any]:
    prompt = f"""
Genera un estado de salud utilizando exclusivamente la siguiente información consolidada.

Institución: {INSTITUTION_NAME}
Fecha de elaboración: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

Información consolidada:
{json.dumps(merged, ensure_ascii=False, indent=2)}
"""
    data = chat_json(SYSTEM_PROMPT_FINAL_ESTADO, prompt)
    return finalize_estado(data, merged)

def build_resumen_clinico(merged: Dict[str, Any]) -> Dict[str, Any]:
    extended_context = build_extended_context(merged)
    prompt = f"""
Genera un resumen clínico documental utilizando exclusivamente la siguiente información consolidada y ampliada.

Institución: {INSTITUTION_NAME}
Fecha de elaboración: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

Información ampliada:
{extended_context}
"""
    data = chat_json(SYSTEM_PROMPT_FINAL_RESUMEN, prompt)
    return finalize_resumen(data, merged)

def build_cronologia(merged: Dict[str, Any]) -> Dict[str, Any]:
    extended_context = build_extended_context(merged)
    prompt = f"""
Genera una cronología médica documental utilizando exclusivamente la siguiente información consolidada y ampliada.

Institución: {INSTITUTION_NAME}
Fecha de elaboración: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

Información ampliada:
{extended_context}
"""
    data = chat_json(SYSTEM_PROMPT_FINAL_CRONOLOGIA, prompt)
    return finalize_cronologia(data, merged)

def build_estado_familiar(merged: Dict[str, Any]) -> Dict[str, Any]:
    prompt = f"""
Genera un estado de salud para familiar utilizando exclusivamente la siguiente información consolidada.

Institución: {INSTITUTION_NAME}
Fecha de elaboración: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

Información consolidada:
{json.dumps(merged, ensure_ascii=False, indent=2)}
"""
    data = chat_json(SYSTEM_PROMPT_FINAL_ESTADO_FAMILIAR, prompt)
    return finalize_estado_familiar(data, merged)

def build_estado_autoridad(merged: Dict[str, Any]) -> Dict[str, Any]:
    extended_context = build_extended_context(merged)
    prompt = f"""
Genera un estado de salud para autoridad utilizando exclusivamente la siguiente información consolidada y ampliada.

Institución: {INSTITUTION_NAME}
Fecha de elaboración: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

Información ampliada:
{extended_context}
"""
    data = chat_json(SYSTEM_PROMPT_FINAL_ESTADO_AUTORIDAD, prompt)
    return finalize_estado_autoridad(data, merged)

def build_estado_institucional(merged: Dict[str, Any]) -> Dict[str, Any]:
    extended_context = build_extended_context(merged)
    prompt = f"""
Genera un estado de salud institucional utilizando exclusivamente la siguiente información consolidada y ampliada.

Institución: {INSTITUTION_NAME}
Fecha de elaboración: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

Información ampliada:
{extended_context}
"""
    data = chat_json(SYSTEM_PROMPT_FINAL_ESTADO_INSTITUCIONAL, prompt)
    return finalize_estado_institucional(data, merged)

def finalize_identificacion(data: Dict[str, Any], merged: Dict[str, Any]) -> None:
    data.setdefault("institucion", INSTITUTION_NAME)
    data.setdefault("identificacion", {})
    data["identificacion"].setdefault("nombre", merged["paciente"]["nombre"])
    data["identificacion"].setdefault("nss", merged["paciente"]["nss"])
    data["identificacion"].setdefault("edad", merged["paciente"]["edad"])
    data["identificacion"].setdefault("sexo", merged["paciente"]["sexo"])
    data["identificacion"]["nombre"] = clean_name(data["identificacion"].get("nombre"))
    data["identificacion"]["nss"] = validate_nss(data["identificacion"].get("nss", ""))
    data.setdefault("fecha_referencia", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    data.setdefault("fuentes", merged.get("fuentes", []))
    data.setdefault("medicos_tratantes", merged.get("medicos_tratantes", []))
    data.setdefault("servicios_clinicos", merged.get("servicios_clinicos", []))
    data.setdefault("signos_vitales", merged.get("signos_vitales", []))

def finalize_estado(data: Dict[str, Any], merged: Dict[str, Any]) -> Dict[str, Any]:
    data.setdefault("tipo_documento", "estado_salud")
    finalize_identificacion(data, merged)
    data.setdefault("resumen_clinico", "No se documenta")
    data.setdefault("diagnosticos_documentados", merged.get("diagnosticos_globales", []))
    data.setdefault("estado_actual", "No se documenta")
    data.setdefault("tratamiento_actual_documentado", merged.get("tratamiento_global_documentado", []))
    data.setdefault("pronostico_documentado", "No se documenta")
    data.setdefault("observaciones", merged.get("observaciones", []))

    if not data.get("texto_final"):
        ident = data["identificacion"]
        data["texto_final"] = (
            f"ESTADO DE SALUD (BORRADOR)\n\n"
            f"Unidad: {data['institucion']}\n"
            f"Fecha de referencia: {data['fecha_referencia']}\n\n"
            f"Nombre: {ident.get('nombre', 'No se documenta')}\n"
            f"NSS: {ident.get('nss', 'No se documenta')}\n"
            f"Edad: {ident.get('edad', 'No se documenta')}\n"
            f"Sexo: {ident.get('sexo', 'No se documenta')}\n\n"
            f"Servicios clínicos documentados:\n{format_list(data.get('servicios_clinicos'))}\n\n"
            f"Médicos tratantes documentados:\n{format_list(data.get('medicos_tratantes'))}\n\n"
            f"Signos vitales documentados:\n{format_list(data.get('signos_vitales'))}\n\n"
            f"Resumen clínico:\n{data.get('resumen_clinico', 'No se documenta')}\n\n"
            f"Diagnósticos documentados:\n{format_list(data.get('diagnosticos_documentados'))}\n\n"
            f"Estado actual documentado:\n{data.get('estado_actual', 'No se documenta')}\n\n"
            f"Tratamiento actual documentado:\n{format_list(data.get('tratamiento_actual_documentado'))}\n\n"
            f"Pronóstico documentado:\n{data.get('pronostico_documentado', 'No se documenta')}\n\n"
            f"Observaciones:\n{format_list(data.get('observaciones'))}\n\n"
            f"Nota: Documento elaborado exclusivamente con base en la información documental analizada; sujeto a validación por personal autorizado."
        )
    return data

def finalize_resumen(data: Dict[str, Any], merged: Dict[str, Any]) -> Dict[str, Any]:
    data.setdefault("tipo_documento", "resumen_clinico")
    finalize_identificacion(data, merged)
    data.setdefault("resumen_clinico", "No se documenta")
    data.setdefault("diagnosticos_documentados", merged.get("diagnosticos_globales", []))
    data.setdefault("tratamiento_actual_documentado", merged.get("tratamiento_global_documentado", []))
    data.setdefault("observaciones", merged.get("observaciones", []))

    if not data.get("texto_final"):
        ident = data["identificacion"]
        data["texto_final"] = (
            f"RESUMEN CLÍNICO DOCUMENTAL\n\n"
            f"Unidad: {data['institucion']}\n"
            f"Fecha de referencia: {data['fecha_referencia']}\n\n"
            f"Nombre: {ident.get('nombre', 'No se documenta')}\n"
            f"NSS: {ident.get('nss', 'No se documenta')}\n"
            f"Edad: {ident.get('edad', 'No se documenta')}\n"
            f"Sexo: {ident.get('sexo', 'No se documenta')}\n\n"
            f"Servicios clínicos documentados:\n{format_list(data.get('servicios_clinicos'))}\n\n"
            f"Médicos tratantes documentados:\n{format_list(data.get('medicos_tratantes'))}\n\n"
            f"Signos vitales documentados:\n{format_list(data.get('signos_vitales'))}\n\n"
            f"Resumen clínico:\n{data.get('resumen_clinico', 'No se documenta')}\n\n"
            f"Diagnósticos documentados:\n{format_list(data.get('diagnosticos_documentados'))}\n\n"
            f"Tratamiento actual documentado:\n{format_list(data.get('tratamiento_actual_documentado'))}\n\n"
            f"Observaciones:\n{format_list(data.get('observaciones'))}"
        )
    return data

def finalize_cronologia(data: Dict[str, Any], merged: Dict[str, Any]) -> Dict[str, Any]:
    data.setdefault("tipo_documento", "cronologia_medica")
    finalize_identificacion(data, merged)
    data.setdefault("cronologia", merged.get("notas_cronologicas", []))

    if not data.get("texto_final"):
        lines = [
            "CRONOLOGÍA MÉDICA DOCUMENTAL",
            "",
            f"Unidad: {data['institucion']}",
            f"Fecha de referencia: {data['fecha_referencia']}",
            "",
        ]

        ident = data["identificacion"]
        lines.extend([
            f"Nombre: {ident.get('nombre', 'No se documenta')}",
            f"NSS: {ident.get('nss', 'No se documenta')}",
            f"Edad: {ident.get('edad', 'No se documenta')}",
            f"Sexo: {ident.get('sexo', 'No se documenta')}",
            "",
            f"Servicios clínicos documentados: {', '.join(data.get('servicios_clinicos', [])) or 'No se documenta'}",
            f"Médicos tratantes documentados: {', '.join(data.get('medicos_tratantes', [])) or 'No se documenta'}",
            "",
        ])

        cronologia = data.get("cronologia", [])
        if not cronologia:
            lines.append("No se documenta cronología.")
        else:
            for i, item in enumerate(cronologia, start=1):
                lines.extend([
                    f"{i}. Fecha: {item.get('fecha', 'No se documenta')}",
                    f"   Servicio: {item.get('servicio', 'No se documenta')}",
                    f"   Tipo de nota: {item.get('tipo_nota', 'No se documenta')}",
                    f"   Médico tratante: {item.get('medico_tratante', 'No se documenta')}",
                    f"   Signos vitales: {', '.join(item.get('signos_vitales', [])) or 'No se documenta'}",
                    f"   Resumen: {item.get('resumen', 'No se documenta')}",
                    f"   Diagnósticos: {', '.join(item.get('diagnosticos', [])) or 'No se documenta'}",
                    f"   Tratamiento: {', '.join(item.get('tratamiento', [])) or 'No se documenta'}",
                    f"   Estado actual: {item.get('estado_actual', 'No se documenta')}",
                    f"   Pronóstico: {item.get('pronostico', 'No se documenta')}",
                    "",
                ])

        data["texto_final"] = "\n".join(lines)
    return data

def finalize_estado_familiar(data: Dict[str, Any], merged: Dict[str, Any]) -> Dict[str, Any]:
    data.setdefault("tipo_documento", "estado_familiar")
    finalize_identificacion(data, merged)
    data.setdefault("resumen_clinico", "No se documenta")
    data.setdefault("estado_actual", "No se documenta")
    data.setdefault("tratamiento_actual_documentado", merged.get("tratamiento_global_documentado", []))
    data.setdefault("observaciones", merged.get("observaciones", []))

    if not data.get("texto_final"):
        ident = data["identificacion"]
        data["texto_final"] = (
            f"ESTADO DE SALUD PARA FAMILIAR (BORRADOR)\n\n"
            f"Unidad: {data['institucion']}\n"
            f"Fecha de referencia: {data['fecha_referencia']}\n\n"
            f"Nombre: {ident.get('nombre', 'No se documenta')}\n"
            f"Edad: {ident.get('edad', 'No se documenta')}\n"
            f"Sexo: {ident.get('sexo', 'No se documenta')}\n\n"
            f"Servicios clínicos documentados:\n{format_list(data.get('servicios_clinicos'))}\n\n"
            f"Resumen de la evolución clínica:\n{data.get('resumen_clinico', 'No se documenta')}\n\n"
            f"Estado actual documentado:\n{data.get('estado_actual', 'No se documenta')}\n\n"
            f"Tratamiento actual documentado:\n{format_list(data.get('tratamiento_actual_documentado'))}\n\n"
            f"Observaciones:\n{format_list(data.get('observaciones'))}\n\n"
            f"Nota: La presente información constituye un borrador informativo elaborado con base en las notas médicas analizadas y no sustituye la información proporcionada directamente por el médico tratante."
        )
    return data

def finalize_estado_autoridad(data: Dict[str, Any], merged: Dict[str, Any]) -> Dict[str, Any]:
    data.setdefault("tipo_documento", "estado_autoridad")
    finalize_identificacion(data, merged)
    data.setdefault("resumen_clinico", "No se documenta")
    data.setdefault("diagnosticos_documentados", merged.get("diagnosticos_globales", []))
    data.setdefault("estado_actual", "No se documenta")
    data.setdefault("tratamiento_actual_documentado", merged.get("tratamiento_global_documentado", []))
    data.setdefault("pronostico_documentado", "No se documenta")
    data.setdefault("observaciones", merged.get("observaciones", []))

    if not data.get("texto_final"):
        ident = data["identificacion"]
        data["texto_final"] = (
            f"ESTADO DE SALUD PARA AUTORIDAD (BORRADOR)\n\n"
            f"Unidad: {data['institucion']}\n"
            f"Fecha de referencia: {data['fecha_referencia']}\n\n"
            f"Nombre: {ident.get('nombre', 'No se documenta')}\n"
            f"NSS: {ident.get('nss', 'No se documenta')}\n"
            f"Edad: {ident.get('edad', 'No se documenta')}\n"
            f"Sexo: {ident.get('sexo', 'No se documenta')}\n\n"
            f"Servicios clínicos documentados:\n{format_list(data.get('servicios_clinicos'))}\n\n"
            f"Médicos tratantes documentados:\n{format_list(data.get('medicos_tratantes'))}\n\n"
            f"Signos vitales documentados:\n{format_list(data.get('signos_vitales'))}\n\n"
            f"Resumen clínico documental:\n{data.get('resumen_clinico', 'No se documenta')}\n\n"
            f"Diagnósticos documentados:\n{format_list(data.get('diagnosticos_documentados'))}\n\n"
            f"Estado actual documentado:\n{data.get('estado_actual', 'No se documenta')}\n\n"
            f"Tratamiento actual documentado:\n{format_list(data.get('tratamiento_actual_documentado'))}\n\n"
            f"Pronóstico documentado:\n{data.get('pronostico_documentado', 'No se documenta')}\n\n"
            f"Observaciones:\n{format_list(data.get('observaciones'))}\n\n"
            f"Nota: Documento elaborado exclusivamente con base en la información documental disponible, sin constituir opinión pericial ni pronunciamiento médico-legal."
        )
    return data

def finalize_estado_institucional(data: Dict[str, Any], merged: Dict[str, Any]) -> Dict[str, Any]:
    data.setdefault("tipo_documento", "estado_institucional")
    finalize_identificacion(data, merged)
    data.setdefault("resumen_clinico", "No se documenta")
    data.setdefault("diagnosticos_documentados", merged.get("diagnosticos_globales", []))
    data.setdefault("estado_actual", "No se documenta")
    data.setdefault("tratamiento_actual_documentado", merged.get("tratamiento_global_documentado", []))
    data.setdefault("pronostico_documentado", "No se documenta")
    data.setdefault("observaciones", merged.get("observaciones", []))

    if not data.get("texto_final"):
        ident = data["identificacion"]
        data["texto_final"] = (
            f"ESTADO DE SALUD INSTITUCIONAL (BORRADOR)\n\n"
            f"Unidad: {data['institucion']}\n"
            f"Fecha de referencia: {data['fecha_referencia']}\n\n"
            f"Nombre: {ident.get('nombre', 'No se documenta')}\n"
            f"NSS: {ident.get('nss', 'No se documenta')}\n"
            f"Edad: {ident.get('edad', 'No se documenta')}\n"
            f"Sexo: {ident.get('sexo', 'No se documenta')}\n\n"
            f"Servicios clínicos documentados:\n{format_list(data.get('servicios_clinicos'))}\n\n"
            f"Médicos tratantes documentados:\n{format_list(data.get('medicos_tratantes'))}\n\n"
            f"Signos vitales documentados:\n{format_list(data.get('signos_vitales'))}\n\n"
            f"Resumen clínico documental:\n{data.get('resumen_clinico', 'No se documenta')}\n\n"
            f"Diagnósticos documentados:\n{format_list(data.get('diagnosticos_documentados'))}\n\n"
            f"Estado actual documentado:\n{data.get('estado_actual', 'No se documenta')}\n\n"
            f"Tratamiento actual documentado:\n{format_list(data.get('tratamiento_actual_documentado'))}\n\n"
            f"Pronóstico documentado:\n{data.get('pronostico_documentado', 'No se documenta')}\n\n"
            f"Observaciones:\n{format_list(data.get('observaciones'))}\n\n"
            f"Nota: Borrador institucional elaborado exclusivamente con base en la revisión documental de las notas médicas aportadas, sujeto a validación por personal autorizado."
        )
    return data

# =========================
# RENDER
# =========================

def render_estado_message(data: Dict[str, Any]) -> str:
    return data.get("texto_final", "No se documenta")

def render_resumen_message(data: Dict[str, Any]) -> str:
    return data.get("texto_final", "No se documenta")

def render_cronologia_message(data: Dict[str, Any]) -> str:
    return data.get("texto_final", "No se documenta")

def render_estado_familiar_message(data: Dict[str, Any]) -> str:
    return data.get("texto_final", "No se documenta")

def render_estado_autoridad_message(data: Dict[str, Any]) -> str:
    return data.get("texto_final", "No se documenta")

def render_estado_institucional_message(data: Dict[str, Any]) -> str:
    return data.get("texto_final", "No se documenta")

def split_message_for_telegram(message: str, max_len: int = TELEGRAM_MAX_MESSAGE) -> List[str]:
    if len(message) <= max_len:
        return [message]

    parts = []
    current = ""

    for paragraph in message.split("\n"):
        candidate = f"{current}\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= max_len:
            current = candidate
        else:
            if current:
                parts.append(current)
            if len(paragraph) <= max_len:
                current = paragraph
            else:
                start = 0
                while start < len(paragraph):
                    end = start + max_len
                    parts.append(paragraph[start:end])
                    start = end
                current = ""

    if current:
        parts.append(current)

    return parts

# =========================
# PDF Y WORD
# =========================

def build_pdf(title: str, body_text: str, output_path: str) -> None:
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="CustomTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        alignment=TA_LEFT,
        spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        name="CustomBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        alignment=TA_LEFT,
        spaceAfter=8,
    ))

    elements = [
        Paragraph(escape_pdf_text(title), styles["CustomTitle"]),
        Spacer(1, 6),
        Paragraph(escape_pdf_text(body_text), styles["CustomBody"]),
    ]

    doc.build(elements)

def build_docx(title: str, body_text: str, output_path: str) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("")

    for line in body_text.split("\n"):
        doc.add_paragraph(line)

    doc.save(output_path)

# =========================
# PIPELINE GENERAL
# =========================

def generate_document_from_text(text: str, mode: str) -> Dict[str, Any]:
    merged = process_medical_text(text)

    if mode == "estado":
        return build_estado_salud(merged)
    if mode == "resumen":
        return build_resumen_clinico(merged)
    if mode == "cronologia":
        return build_cronologia(merged)
    if mode == "estado_familiar":
        return build_estado_familiar(merged)
    if mode == "estado_autoridad":
        return build_estado_autoridad(merged)
    if mode == "estado_institucional":
        return build_estado_institucional(merged)

    raise ValueError(f"Modo no soportado: {mode}")

def process_pdf_and_generate(pdf_bytes: bytes, mode: str) -> Dict[str, Any]:
    text = extract_text_from_pdf_bytes(pdf_bytes)

    if not text.strip():
        return {"_error": "No se pudo extraer texto del PDF."}

    if len(text.strip()) < MIN_TEXT_LENGTH:
        return {"_error": "El PDF no contiene suficiente texto legible para generar un resultado confiable."}

    data = generate_document_from_text(text, mode)
    return {"data": data}

async def send_outputs(
    update: Update,
    data: Dict[str, Any],
    message: str,
    base_filename: str,
    context: ContextTypes.DEFAULT_TYPE
) -> None:
    temp_paths = []

    try:
        for chunk in split_message_for_telegram(message):
            msg = await update.message.reply_text(chunk)
            track_bot_message_id(context, msg.message_id)

        title_map = {
            "estado_salud": "ESTADO DE SALUD",
            "resumen_clinico": "RESUMEN CLÍNICO DOCUMENTAL",
            "cronologia_medica": "CRONOLOGÍA MÉDICA DOCUMENTAL",
            "estado_familiar": "ESTADO DE SALUD PARA FAMILIAR",
            "estado_autoridad": "ESTADO DE SALUD PARA AUTORIDAD",
            "estado_institucional": "ESTADO DE SALUD INSTITUCIONAL",
        }

        title = title_map.get(data.get("tipo_documento", ""), "DOCUMENTO")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename_base = f"{base_filename}_{timestamp}"

        if INCLUDE_PDF_FILE:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as f:
                pdf_path = f.name
                temp_paths.append(pdf_path)

            build_pdf(
                title=title,
                body_text=data.get("texto_final", "No se documenta"),
                output_path=pdf_path,
            )

            with open(pdf_path, "rb") as f:
                msg = await update.message.reply_document(
                    document=f,
                    filename=f"{filename_base}.pdf",
                    caption="Archivo PDF generado"
                )
                track_bot_message_id(context, msg.message_id)

        if INCLUDE_DOCX_FILE:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
                docx_path = f.name
                temp_paths.append(docx_path)

            build_docx(
                title=title,
                body_text=data.get("texto_final", "No se documenta"),
                output_path=docx_path,
            )

            with open(docx_path, "rb") as f:
                msg = await update.message.reply_document(
                    document=f,
                    filename=f"{filename_base}.docx",
                    caption="Archivo Word generado"
                )
                track_bot_message_id(context, msg.message_id)

        finish_text = "✅ Proceso concluido.\n\nSelecciona una opción:"
        msg = await update.message.reply_text(
            finish_text,
            reply_markup=build_finish_keyboard()
        )
        track_bot_message_id(context, msg.message_id)

        context.user_data["mode"] = "estado"

    finally:
        for path in temp_paths:
            try:
                if os.path.exists(path):
                    os.remove(path)
            except Exception:
                logger.warning("No se pudo eliminar archivo temporal: %s", path)

# =========================
# TELEGRAM
# =========================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    context.user_data["mode"] = "estado"
    await send_tracked_text(update, context, HELP_TEXT)

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await send_tracked_text(update, context, HELP_TEXT)

async def set_estado(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    context.user_data["mode"] = "estado"
    await send_tracked_text(update, context, "Modo seleccionado: estado de salud. Ahora envía el PDF.")

async def set_resumen(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    context.user_data["mode"] = "resumen"
    await send_tracked_text(update, context, "Modo seleccionado: resumen clínico. Ahora envía el PDF.")

async def set_cronologia(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    context.user_data["mode"] = "cronologia"
    await send_tracked_text(update, context, "Modo seleccionado: cronología médica. Ahora envía el PDF.")

async def set_estado_familiar(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    context.user_data["mode"] = "estado_familiar"
    await send_tracked_text(update, context, "Modo seleccionado: estado para familiar. Ahora envía el PDF.")

async def set_estado_autoridad(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    context.user_data["mode"] = "estado_autoridad"
    await send_tracked_text(update, context, "Modo seleccionado: estado para autoridad. Ahora envía el PDF.")

async def set_estado_institucional(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    context.user_data["mode"] = "estado_institucional"
    await send_tracked_text(update, context, "Modo seleccionado: estado institucional. Ahora envía el PDF.")

async def handle_inline_actions(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    query = update.callback_query
    await query.answer()

    action = query.data

    if action == "nuevo_reporte":
        context.user_data["mode"] = "estado"
        await send_tracked_text_from_query(
            query,
            context,
            "Modo reiniciado en /estado.\n\nEnvía un nuevo PDF o selecciona otro modo con /help."
        )
        return

    if action == "ver_menu":
        await send_tracked_text_from_query(query, context, HELP_TEXT)
        return

    if action == "limpiar_bot":
        chat_id = query.message.chat_id
        bot_message_ids = context.user_data.get("bot_message_ids", [])

        deleted = 0
        for msg_id in reversed(bot_message_ids[-80:]):
            try:
                await context.bot.delete_message(chat_id=chat_id, message_id=msg_id)
                deleted += 1
            except Exception:
                pass

        context.user_data["bot_message_ids"] = []

        try:
            msg = await query.message.reply_text(
                f"Se intentaron limpiar {deleted} mensajes del bot.\n\n{HELP_TEXT}"
            )
            track_bot_message_id(context, msg.message_id)
        except Exception:
            pass

async def handle_pdf(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    document = update.message.document
    if not document:
        return

    mode = context.user_data.get("mode", "estado")

    logger.info(
        "Usuario=%s Archivo=%s Modo=%s Tamaño=%s",
        update.effective_user.id if update.effective_user else "desconocido",
        getattr(document, "file_name", "sin_nombre"),
        mode,
        getattr(document, "file_size", "sin_dato"),
    )

    if document.file_size and document.file_size > MAX_FILE_SIZE_MB * 1024 * 1024:
        await send_tracked_text(
            update,
            context,
            f"El archivo excede el tamaño máximo permitido de {MAX_FILE_SIZE_MB} MB."
        )
        return

    await send_tracked_text(update, context, f"Procesando PDF en modo: {mode}")

    try:
        tg_file = await document.get_file()
        pdf_bytes = await tg_file.download_as_bytearray()

        result = await asyncio.to_thread(
            process_pdf_and_generate,
            bytes(pdf_bytes),
            mode
        )

        if "_error" in result:
            await send_tracked_text(update, context, result["_error"])
            return

        data = result["data"]

        if mode == "estado":
            message = render_estado_message(data)
            base_filename = "estado_salud"
        elif mode == "resumen":
            message = render_resumen_message(data)
            base_filename = "resumen_clinico"
        elif mode == "cronologia":
            message = render_cronologia_message(data)
            base_filename = "cronologia_medica"
        elif mode == "estado_familiar":
            message = render_estado_familiar_message(data)
            base_filename = "estado_familiar"
        elif mode == "estado_autoridad":
            message = render_estado_autoridad_message(data)
            base_filename = "estado_autoridad"
        else:
            message = render_estado_institucional_message(data)
            base_filename = "estado_institucional"

        await send_outputs(update, data, message, base_filename, context)

    except Exception as e:
        logger.exception("Error al procesar el PDF")
        error_text = str(e).lower()

        if "429" in error_text or "quota" in error_text or "rate limit" in error_text:
            await send_tracked_text(
                update,
                context,
                "No fue posible procesar el documento por límite de cuota o velocidad del servicio de IA. Intenta nuevamente en unos minutos."
            )
        else:
            await send_tracked_text(
                update,
                context,
                f"Ocurrió un error al procesar el PDF: {e}"
            )

def main() -> None:
    app = Application.builder().token(TELEGRAM_BOT_TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("help", help_command))
    app.add_handler(CommandHandler("estado", set_estado))
    app.add_handler(CommandHandler("resumen", set_resumen))
    app.add_handler(CommandHandler("cronologia", set_cronologia))
    app.add_handler(CommandHandler("estado_familiar", set_estado_familiar))
    app.add_handler(CommandHandler("estado_autoridad", set_estado_autoridad))
    app.add_handler(CommandHandler("estado_institucional", set_estado_institucional))
    app.add_handler(CallbackQueryHandler(handle_inline_actions))
    app.add_handler(MessageHandler(filters.Document.PDF, handle_pdf))

    app.run_polling()

if __name__ == "__main__":
    main()
